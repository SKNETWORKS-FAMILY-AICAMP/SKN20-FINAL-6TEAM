"""액션 실행기 모듈.

문서 생성 등의 액션을 실행하는 Action Executor를 구현합니다.
"""

import base64
import logging
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any
from xml.sax.saxutils import escape as xml_escape

import httpx

from schemas.request import ContractRequest
from schemas.response import DocumentResponse
from utils.config import get_settings
from utils.s3_client import get_s3_client

logger = logging.getLogger(__name__)

# 한글 폰트 경로 후보 (시스템 → 프로젝트 내장 순)
_KOREAN_FONT_PATHS = [
    # Docker (apt install fonts-nanum)
    Path("/usr/share/fonts/truetype/nanum"),
    # 프로젝트 내장 폴백
    Path(__file__).resolve().parent.parent / "fonts",
]

_font_registered = False


def _register_korean_font() -> None:
    """한글 폰트를 reportlab에 등록합니다 (최초 1회)."""
    global _font_registered
    if _font_registered:
        return

    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    for font_dir in _KOREAN_FONT_PATHS:
        regular = font_dir / "NanumGothic.ttf"
        bold = font_dir / "NanumGothicBold.ttf"
        if regular.exists() and bold.exists():
            pdfmetrics.registerFont(TTFont("KoreanFont", str(regular)))
            pdfmetrics.registerFont(TTFont("KoreanFont-Bold", str(bold)))
            pdfmetrics.registerFontFamily("KoreanFont", normal="KoreanFont", bold="KoreanFont-Bold")
            _font_registered = True
            logger.info("한글 폰트 등록 완료: %s", font_dir)
            return

    logger.warning("한글 폰트를 찾을 수 없습니다. PDF에서 한글이 깨질 수 있습니다.")


class ActionExecutor:
    """액션 실행기 클래스.

    문서 생성 등의 액션을 실행합니다.

    지원 문서:
    - 근로계약서 (PDF)
    - 취업규칙 템플릿 (DOCX)
    - 사업계획서 템플릿 (DOCX)

    Attributes:
        settings: 설정 객체
        output_dir: 출력 디렉토리

    Example:
        >>> executor = ActionExecutor()
        >>> result = executor.generate_labor_contract(contract_data)
        >>> print(result.file_name)
    """

    def __init__(self) -> None:
        """ActionExecutor를 초기화합니다."""
        self.settings = get_settings()
        self.output_dir = Path(__file__).parent.parent / "output"
        self.output_dir.mkdir(exist_ok=True)

    # ---------- S3 업로드 + DB 저장 ----------

    def _upload_and_save(
        self,
        response: DocumentResponse,
        user_id: int | None,
        company_id: int | None = None,
    ) -> DocumentResponse:
        """생성된 문서를 S3에 업로드하고 DB에 메타데이터를 저장합니다.

        user_id가 None이면 S3/DB 저장을 스킵합니다 (게스트 사용자).

        Args:
            response: 문서 생성 응답 (file_content base64 포함)
            user_id: 사용자 ID (None이면 스킵)
            company_id: 회사 ID (선택)

        Returns:
            s3_key, file_id가 추가된 DocumentResponse
        """
        if not user_id or not response.success or not response.file_content:
            return response

        try:
            s3_client = get_s3_client()
            file_bytes = base64.b64decode(response.file_content)
            ext = (response.file_name or "doc").rsplit(".", 1)[-1]
            s3_key = s3_client._generate_key(user_id, response.doc_type_id, ext)

            content_type = "application/pdf" if ext == "pdf" else (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            s3_client.upload_document(file_bytes, s3_key, content_type)

            file_id = self._save_file_metadata(
                user_id=user_id,
                company_id=company_id,
                doc_type_id=response.doc_type_id,
                file_name=response.file_name or "",
                s3_key=s3_key,
                file_size=len(file_bytes),
                file_format=ext,
            )

            response.s3_key = s3_key
            response.file_id = file_id
            logger.info("문서 S3 업로드 + DB 저장 완료: s3_key=%s file_id=%s", s3_key, file_id)
        except Exception:
            logger.error("문서 S3/DB 저장 실패 (문서 생성은 성공)", exc_info=True)

        return response

    def _save_file_metadata(
        self,
        user_id: int,
        company_id: int | None,
        doc_type_id: str,
        file_name: str,
        s3_key: str,
        file_size: int,
        file_format: str,
        parent_file_id: int | None = None,
    ) -> int | None:
        """Backend API를 호출하여 파일 메타데이터를 DB에 저장합니다.

        Args:
            user_id: 사용자 ID
            company_id: 회사 ID
            doc_type_id: 문서 유형
            file_name: 파일명
            s3_key: S3 오브젝트 키
            file_size: 파일 크기 (bytes)
            file_format: 파일 형식 (pdf, docx)
            parent_file_id: 수정 원본 파일 ID (버전관리)

        Returns:
            생성된 file_id 또는 None (실패 시)
        """
        settings = get_settings()
        url = f"{settings.backend_internal_url}/documents/save"
        payload = {
            "user_id": user_id,
            "company_id": company_id,
            "doc_type_id": doc_type_id,
            "file_name": file_name,
            "s3_key": s3_key,
            "file_size": file_size,
            "file_format": file_format,
        }
        if parent_file_id:
            payload["parent_file_id"] = parent_file_id

        headers: dict[str, str] = {}
        if settings.rag_api_key:
            headers["X-API-Key"] = settings.rag_api_key

        try:
            with httpx.Client(timeout=10.0) as client:
                resp = client.post(url, json=payload, headers=headers)
            if resp.status_code == 201:
                return resp.json().get("file_id")
            logger.error("DB 메타데이터 저장 실패: status=%d body=%s", resp.status_code, resp.text[:500])
        except Exception:
            logger.error("Backend API 호출 실패", exc_info=True)
        return None

    def generate_labor_contract(
        self,
        request: ContractRequest,
        user_id: int | None = None,
        company_id: int | None = None,
    ) -> DocumentResponse:
        """근로계약서를 생성합니다.

        Args:
            request: 근로계약서 생성 요청
            user_id: 사용자 ID (있으면 S3/DB 저장)
            company_id: 회사 ID (선택)

        Returns:
            문서 생성 응답
        """
        try:
            if request.format == "pdf":
                result = self._generate_contract_pdf(request)
            elif request.format == "docx":
                result = self._generate_contract_docx(request)
            else:
                return DocumentResponse(
                    success=False,
                    doc_type_id="labor_contract",
                    message=f"지원하지 않는 형식: {request.format} (pdf 또는 docx만 가능)",
                )
            return self._upload_and_save(result, user_id, company_id)
        except Exception as e:
            return DocumentResponse(
                success=False,
                doc_type_id="labor_contract",
                message=f"근로계약서 생성 실패: {str(e)}",
            )

    def _generate_contract_pdf(
        self,
        request: ContractRequest,
    ) -> DocumentResponse:
        """근로계약서 PDF를 생성합니다."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        _register_korean_font()

        # 버퍼 생성
        buffer = BytesIO()

        # PDF 문서 생성
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        # 스타일 설정
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "Title",
            parent=styles["Title"],
            fontName="KoreanFont-Bold",
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # 가운데 정렬
        )
        normal_style = ParagraphStyle(
            "Normal",
            parent=styles["Normal"],
            fontName="KoreanFont",
            fontSize=11,
            leading=16,
        )

        # 내용 구성
        elements = []

        # 제목
        elements.append(Paragraph("표준 근로계약서", title_style))
        elements.append(Spacer(1, 20))

        # 계약 당사자 (사용자 입력 이스케이프)
        company_name = "회사명 미기재"
        if request.company_context:
            company_name = request.company_context.company_name or company_name

        elements.append(Paragraph(
            f"<b>사업주(갑)</b>: {xml_escape(company_name)}",
            normal_style
        ))
        elements.append(Paragraph(
            f"<b>근로자(을)</b>: {xml_escape(request.employee_name)}",
            normal_style
        ))
        elements.append(Spacer(1, 20))

        # 계약 내용
        contract_type = "무기계약" if request.is_permanent else "기간제 계약"
        contract_period = "정함 없음"
        if not request.is_permanent and request.contract_end_date:
            contract_period = f"{request.contract_start_date} ~ {request.contract_end_date}"

        # 사용자 입력 이스케이프 (ReportLab HTML 인젝션 방지)
        esc = xml_escape

        # 제7조 (임금) — 상세 항목 조립
        salary_lines = [
            f"기본급: 월 {request.base_salary:,}원",
            f"연장근로수당: 통상임금의 {request.overtime_pay_rate}%",
            f"야간근로수당(22시~06시): 통상임금의 {request.night_pay_rate}%",
            f"휴일근로수당: 통상임금의 {request.holiday_pay_rate}%",
        ]
        if request.bonus:
            salary_lines.append(f"상여금: {esc(request.bonus)}")
        if request.allowances:
            salary_lines.append(f"제 수당: {esc(request.allowances)}")
        salary_lines.append(f"지급 방법: {esc(request.payment_method)}")
        salary_lines.append(f"지급일: 매월 {request.payment_date}일")
        salary_block = "<br/>".join(salary_lines)

        # 제9조 (단시간근로자 특약) — 조건부
        part_time_block = ""
        if request.is_part_time:
            wh = request.weekly_work_hours or 0
            part_time_block = f"""
        <b>제9조 (단시간근로자 특약)</b><br/>
        본 계약의 근로자는 단시간근로자(근로기준법 제2조 제1항 제9호)에 해당합니다.<br/>
        주 소정근로시간: {wh}시간<br/>
        연차유급휴가, 퇴직금 등 법정 급여는 통상 근로자의 소정근로시간에 비례하여 산정합니다 (근로기준법 제18조).<br/><br/>
"""

        # 기타 조항 번호 (단시간근로자 유무에 따라 동적)
        other_num = "제10조" if request.is_part_time else "제9조"

        content = f"""
        <b>제1조 (계약 유형)</b><br/>
        본 계약은 {esc(contract_type)}입니다.<br/><br/>

        <b>제2조 (계약 기간)</b><br/>
        계약 시작일: {esc(request.contract_start_date)}<br/>
        계약 기간: {esc(contract_period)}<br/><br/>

        <b>제3조 (근무 장소)</b><br/>
        {esc(request.workplace)}<br/><br/>

        <b>제4조 (업무 내용)</b><br/>
        직위: {esc(request.job_title)}<br/>
        담당 업무: {esc(request.job_description)}<br/><br/>

        <b>제5조 (근무 시간)</b><br/>
        근무 시간: {esc(request.work_start_time)} ~ {esc(request.work_end_time)}<br/>
        휴게 시간: {esc(request.rest_time)}<br/>
        근무 요일: {esc(request.work_days)}<br/><br/>

        <b>제6조 (휴일)</b><br/>
        {esc(request.holidays)}<br/><br/>

        <b>제7조 (임금)</b><br/>
        {salary_block}<br/><br/>

        <b>제8조 (연차유급휴가)</b><br/>
        연차유급휴가: 연 {request.annual_leave_days}일<br/>
        근로기준법 제60조에서 정한 기준에 따라 산정·부여합니다.<br/><br/>
        """ + part_time_block + f"""
        <b>{other_num} (기타)</b><br/>
        본 계약서에 명시되지 않은 사항은 근로기준법 및 관계 법령에 따릅니다.
        """

        elements.append(Paragraph(content, normal_style))
        elements.append(Spacer(1, 40))

        # 서명란
        today = datetime.now().strftime("%Y년 %m월 %d일")
        elements.append(Paragraph(f"작성일: {today}", normal_style))
        elements.append(Spacer(1, 20))
        elements.append(Paragraph("사업주(갑):                    (인)", normal_style))
        elements.append(Spacer(1, 10))
        elements.append(Paragraph("근로자(을):                    (인)", normal_style))

        # PDF 빌드
        doc.build(elements)

        # 파일 저장
        file_name = self._build_file_name("labor_contract", "pdf", company_name)
        file_path = self.output_dir / file_name

        buffer.seek(0)
        with open(file_path, "wb") as f:
            f.write(buffer.read())

        # base64 인코딩
        buffer.seek(0)
        file_content = base64.b64encode(buffer.read()).decode("utf-8")

        return DocumentResponse(
            success=True,
            doc_type_id="labor_contract",
            file_path=str(file_path),
            file_name=file_name,
            file_content=file_content,
            message="근로계약서가 생성되었습니다.",
        )

    def _generate_contract_docx(
        self,
        request: ContractRequest,
    ) -> DocumentResponse:
        """근로계약서 DOCX를 생성합니다."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 문서 생성
        doc = Document()

        # 제목
        title = doc.add_heading("표준 근로계약서", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 계약 당사자
        company_name = "회사명 미기재"
        if request.company_context:
            company_name = request.company_context.company_name or company_name

        doc.add_paragraph(f"사업주(갑): {company_name}")
        doc.add_paragraph(f"근로자(을): {request.employee_name}")
        doc.add_paragraph()

        # 계약 내용
        contract_type = "무기계약" if request.is_permanent else "기간제 계약"
        contract_period = "정함 없음"
        if not request.is_permanent and request.contract_end_date:
            contract_period = f"{request.contract_start_date} ~ {request.contract_end_date}"

        doc.add_heading("제1조 (계약 유형)", level=2)
        doc.add_paragraph(f"본 계약은 {contract_type}입니다.")

        doc.add_heading("제2조 (계약 기간)", level=2)
        doc.add_paragraph(f"계약 시작일: {request.contract_start_date}")
        doc.add_paragraph(f"계약 기간: {contract_period}")

        doc.add_heading("제3조 (근무 장소)", level=2)
        doc.add_paragraph(request.workplace)

        doc.add_heading("제4조 (업무 내용)", level=2)
        doc.add_paragraph(f"직위: {request.job_title}")
        doc.add_paragraph(f"담당 업무: {request.job_description}")

        doc.add_heading("제5조 (근무 시간)", level=2)
        doc.add_paragraph(f"근무 시간: {request.work_start_time} ~ {request.work_end_time}")
        doc.add_paragraph(f"휴게 시간: {request.rest_time}")
        doc.add_paragraph(f"근무 요일: {request.work_days}")

        doc.add_heading("제6조 (휴일)", level=2)
        doc.add_paragraph(request.holidays)

        doc.add_heading("제7조 (임금)", level=2)
        doc.add_paragraph(f"기본급: 월 {request.base_salary:,}원")
        doc.add_paragraph(f"연장근로수당: 통상임금의 {request.overtime_pay_rate}%")
        doc.add_paragraph(f"야간근로수당(22시~06시): 통상임금의 {request.night_pay_rate}%")
        doc.add_paragraph(f"휴일근로수당: 통상임금의 {request.holiday_pay_rate}%")
        if request.bonus:
            doc.add_paragraph(f"상여금: {request.bonus}")
        if request.allowances:
            doc.add_paragraph(f"제 수당: {request.allowances}")
        doc.add_paragraph(f"지급 방법: {request.payment_method}")
        doc.add_paragraph(f"지급일: 매월 {request.payment_date}일")

        doc.add_heading("제8조 (연차유급휴가)", level=2)
        doc.add_paragraph(f"연차유급휴가: 연 {request.annual_leave_days}일")
        doc.add_paragraph("근로기준법 제60조에서 정한 기준에 따라 산정·부여합니다.")

        if request.is_part_time:
            wh = request.weekly_work_hours or 0
            doc.add_heading("제9조 (단시간근로자 특약)", level=2)
            doc.add_paragraph("본 계약의 근로자는 단시간근로자(근로기준법 제2조 제1항 제9호)에 해당합니다.")
            doc.add_paragraph(f"주 소정근로시간: {wh}시간")
            doc.add_paragraph("연차유급휴가, 퇴직금 등 법정 급여는 통상 근로자의 소정근로시간에 비례하여 산정합니다 (근로기준법 제18조).")

        other_num = "제10조" if request.is_part_time else "제9조"
        doc.add_heading(f"{other_num} (기타)", level=2)
        doc.add_paragraph("본 계약서에 명시되지 않은 사항은 근로기준법 및 관계 법령에 따릅니다.")

        # 서명란
        doc.add_paragraph()
        today = datetime.now().strftime("%Y년 %m월 %d일")
        doc.add_paragraph(f"작성일: {today}")
        doc.add_paragraph()
        doc.add_paragraph("사업주(갑):                    (인)")
        doc.add_paragraph("근로자(을):                    (인)")

        # 파일 저장
        file_name = self._build_file_name("labor_contract", "docx", company_name)
        file_path = self.output_dir / file_name

        doc.save(file_path)

        # base64 인코딩
        with open(file_path, "rb") as f:
            file_content = base64.b64encode(f.read()).decode("utf-8")

        return DocumentResponse(
            success=True,
            doc_type_id="labor_contract",
            file_path=str(file_path),
            file_name=file_name,
            file_content=file_content,
            message="근로계약서가 생성되었습니다.",
        )

    def generate_business_plan_template(
        self,
        format: str = "docx",
        user_id: int | None = None,
        company_id: int | None = None,
    ) -> DocumentResponse:
        """사업계획서 템플릿을 생성합니다.

        Args:
            format: 출력 형식 (docx)
            user_id: 사용자 ID (있으면 S3/DB 저장)
            company_id: 회사 ID (선택)

        Returns:
            문서 생성 응답
        """
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # 문서 생성
            doc = Document()

            # 제목
            title = doc.add_heading("사업계획서", level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 목차
            sections = [
                ("1. 사업 개요", [
                    "1.1 사업 아이템 소개",
                    "1.2 창업 동기 및 비전",
                    "1.3 사업 목표",
                ]),
                ("2. 시장 분석", [
                    "2.1 시장 규모 및 성장성",
                    "2.2 경쟁 현황",
                    "2.3 타겟 고객",
                ]),
                ("3. 제품/서비스", [
                    "3.1 제품/서비스 설명",
                    "3.2 차별화 포인트",
                    "3.3 가격 정책",
                ]),
                ("4. 마케팅 전략", [
                    "4.1 마케팅 목표",
                    "4.2 홍보 채널",
                    "4.3 고객 확보 전략",
                ]),
                ("5. 운영 계획", [
                    "5.1 조직 구성",
                    "5.2 인력 계획",
                    "5.3 운영 프로세스",
                ]),
                ("6. 재무 계획", [
                    "6.1 소요 자금",
                    "6.2 매출 계획",
                    "6.3 손익 계획",
                ]),
            ]

            for section_title, subsections in sections:
                doc.add_heading(section_title, level=1)
                for subsection in subsections:
                    doc.add_heading(subsection, level=2)
                    doc.add_paragraph("[내용을 작성하세요]")
                    doc.add_paragraph()

            # 파일 저장
            file_name = self._build_file_name("business_plan", "docx")
            file_path = self.output_dir / file_name

            doc.save(file_path)

            # base64 인코딩
            with open(file_path, "rb") as f:
                file_content = base64.b64encode(f.read()).decode("utf-8")

            result = DocumentResponse(
                success=True,
                doc_type_id="business_plan",
                file_path=str(file_path),
                file_name=file_name,
                file_content=file_content,
                message="사업계획서 템플릿이 생성되었습니다.",
            )
            return self._upload_and_save(result, user_id, company_id)

        except Exception as e:
            return DocumentResponse(
                success=False,
                doc_type_id="business_plan",
                message=f"사업계획서 템플릿 생성 실패: {str(e)}",
            )

    # ---------- 범용 문서 생성 (LLM 기반) ----------

    def generate_document(
        self,
        document_type: str,
        params: dict[str, Any],
        format: str = "docx",
        user_id: int | None = None,
        company_id: int | None = None,
    ) -> DocumentResponse:
        """범용 문서 생성 엔트리포인트.

        Args:
            document_type: 문서 유형 키
            params: 문서 필드 값
            format: 출력 형식 (pdf, docx)
            user_id: 사용자 ID (있으면 S3/DB 저장)
            company_id: 회사 ID (선택)

        Returns:
            문서 생성 응답
        """
        # application_form은 동적 필드이므로 레지스트리 우회
        if document_type == "application_form":
            result = self._generate_application_form(params, format)
            return self._upload_and_save(result, user_id, company_id)

        from agents.document_registry import DOCUMENT_TYPE_REGISTRY

        type_def = DOCUMENT_TYPE_REGISTRY.get(document_type)
        if not type_def:
            return DocumentResponse(
                success=False,
                doc_type_id=document_type,
                message=f"알 수 없는 문서 유형: {document_type}",
            )

        # 하드코딩 문서 → 기존 로직 위임
        if type_def.generation_method == "hardcoded":
            if document_type == "labor_contract":
                return self.generate_labor_contract(
                    ContractRequest(**params), user_id=user_id, company_id=company_id,
                )

        # LLM 기반 생성
        result = self._generate_document_by_llm(type_def, params, format)
        return self._upload_and_save(result, user_id, company_id)

    def _generate_document_by_llm(
        self,
        type_def: Any,
        params: dict[str, Any],
        format: str,
    ) -> DocumentResponse:
        """LLM으로 문서 내용을 생성한 뒤 DOCX/PDF로 변환합니다.

        Args:
            type_def: DocumentTypeDef 객체
            params: 문서 필드 값
            format: 출력 형식

        Returns:
            문서 생성 응답
        """
        from utils import prompts
        from utils.config.llm import create_llm

        try:
            prompt_template = getattr(prompts, type_def.llm_prompt_key, None)
            if not prompt_template:
                return DocumentResponse(
                    success=False,
                    doc_type_id=type_def.type_key,
                    message=f"프롬프트를 찾을 수 없습니다: {type_def.llm_prompt_key}",
                )

            today = datetime.now().strftime("%Y년 %m월 %d일")
            field_text = "\n".join(
                f"- {k}: {v}" for k, v in params.items() if v
            )
            field_text += f"\n- 작성일: {today}"
            full_prompt = prompt_template.format(field_values=field_text)

            token_limit = 8192 if type_def.type_key == "business_plan" else 4096
            llm = create_llm(label="문서생성", temperature=0.3, max_tokens=token_limit)
            response = llm.invoke([{"role": "user", "content": full_prompt}])
            content = response.content

            company_name = params.get("company_name") or params.get("party_a") or params.get("client")
            if format == "pdf":
                return self._build_pdf_from_text(content, type_def.type_key, type_def.label, company_name)
            else:
                return self._build_docx_from_text(content, type_def.type_key, type_def.label, company_name)

        except Exception as e:
            logger.error("LLM 문서 생성 실패 (%s): %s", type_def.type_key, e, exc_info=True)
            return DocumentResponse(
                success=False,
                doc_type_id=type_def.type_key,
                message=f"문서 생성 실패: {str(e)}",
            )

    def _generate_application_form(
        self,
        params: dict[str, Any],
        format: str = "docx",
    ) -> DocumentResponse:
        """S3 신청 양식을 기반으로 사용자 입력값을 채운 문서를 생성합니다.

        Args:
            params: 사용자 입력 필드값 + _form_key, _form_title 메타 키
            format: 출력 형식 (pdf, docx)

        Returns:
            문서 생성 응답
        """
        from utils import prompts
        from utils.config.llm import create_llm
        from utils.file_parser import extract_text_from_base64

        form_key: str = params.pop("_form_key", "")
        form_title: str = params.pop("_form_title", "신청서")

        if not form_key:
            return DocumentResponse(
                success=False,
                doc_type_id="application_form",
                message="양식 키(_form_key)가 전달되지 않았습니다.",
            )

        try:
            # S3에서 원본 양식 다운로드 → 텍스트 추출
            s3 = get_s3_client()
            file_bytes: bytes = s3.get_application_form(form_key)
            file_name: str = form_key.split("/")[-1]
            b64_content: str = base64.b64encode(file_bytes).decode()
            form_text: str = extract_text_from_base64(b64_content, file_name)

            # 사용자 입력값 포매팅
            today: str = datetime.now().strftime("%Y년 %m월 %d일")
            field_text: str = "\n".join(
                f"- {k}: {v}" for k, v in params.items() if v
            )
            field_text += f"\n- 작성일: {today}"

            full_prompt: str = prompts.APPLICATION_FORM_GENERATION_PROMPT.format(
                form_text=form_text,
                field_values=field_text,
            )

            llm = create_llm(label="신청서생성", temperature=0.3, max_tokens=4096)
            response = llm.invoke([{"role": "user", "content": full_prompt}])
            content: str = response.content

            if format == "pdf":
                return self._build_pdf_from_text(content, "application_form", form_title)
            return self._build_docx_from_text(content, "application_form", form_title)

        except Exception as e:
            logger.error("신청서 생성 실패 (%s): %s", form_key, e, exc_info=True)
            return DocumentResponse(
                success=False,
                doc_type_id="application_form",
                message=f"신청서 생성 실패: {str(e)}",
            )

    def modify_document(
        self,
        file_content: str,
        file_name: str,
        instructions: str,
        format: str = "docx",
        user_id: int | None = None,
        document_id: int | None = None,
        company_id: int | None = None,
    ) -> DocumentResponse:
        """업로드된 문서를 LLM으로 수정합니다.

        Args:
            file_content: base64 인코딩된 원본 파일
            file_name: 원본 파일명 (확장자 포함)
            instructions: 수정 지시사항
            format: 출력 형식 (pdf, docx)
            user_id: 사용자 ID (있으면 S3/DB 저장)
            document_id: 원본 문서 ID (버전관리)
            company_id: 회사 ID (선택)

        Returns:
            수정된 문서 응답
        """
        from utils import prompts
        from utils.config.llm import create_llm
        from utils.file_parser import extract_text_from_base64

        try:
            original_text = extract_text_from_base64(file_content, file_name)

            # 비즈니스 문서 적합성 검증
            relevance = self._check_document_relevance(original_text)
            if not relevance.get("is_relevant", True):
                reason = relevance.get("reason", "비즈니스 관련 문서가 아닙니다.")
                return DocumentResponse(
                    success=False,
                    doc_type_id="rejected",
                    message=f"문서 수정이 거부되었습니다: {reason}",
                )

            full_prompt = prompts.DOCUMENT_MODIFY_PROMPT.format(
                original_text=original_text,
                instructions=instructions,
            )

            llm = create_llm(label="문서수정", temperature=0.2, max_tokens=8192)
            response = llm.invoke([{"role": "user", "content": full_prompt}])
            content = response.content

            title = "수정된 문서"
            type_key = "modified_document"

            if format == "pdf":
                result = self._build_pdf_from_text(content, type_key, title)
            else:
                result = self._build_docx_from_text(content, type_key, title)

            # S3 + DB 저장 (수정 버전)
            if user_id and result.success:
                if document_id:
                    # 버전관리: parent_file_id를 통해 원본 연결
                    try:
                        file_bytes = base64.b64decode(result.file_content or "")
                        ext = (result.file_name or "doc").rsplit(".", 1)[-1]
                        s3_client = get_s3_client()
                        s3_key = s3_client._generate_key(user_id, type_key, ext)
                        content_type = "application/pdf" if ext == "pdf" else (
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        s3_client.upload_document(file_bytes, s3_key, content_type)
                        file_id = self._save_file_metadata(
                            user_id=user_id,
                            company_id=company_id,
                            doc_type_id=type_key,
                            file_name=result.file_name or "",
                            s3_key=s3_key,
                            file_size=len(file_bytes),
                            file_format=ext,
                            parent_file_id=document_id,
                        )
                        result.s3_key = s3_key
                        result.file_id = file_id
                    except Exception:
                        logger.error("수정 문서 S3/DB 저장 실패", exc_info=True)
                else:
                    result = self._upload_and_save(result, user_id, company_id)

            return result

        except ValueError as e:
            logger.warning("문서 수정 입력 오류: %s", e)
            return DocumentResponse(
                success=False,
                doc_type_id="modified_document",
                message=str(e),
            )
        except Exception as e:
            logger.error("문서 수정 실패: %s", e, exc_info=True)
            return DocumentResponse(
                success=False,
                doc_type_id="modified_document",
                message=f"문서 수정 실패: {str(e)}",
            )

    def _check_document_relevance(self, text: str) -> dict[str, Any]:
        """문서가 비즈니스 관련 문서인지 LLM으로 검증합니다.

        Args:
            text: 문서에서 추출한 텍스트

        Returns:
            {"is_relevant": bool, "category": str, "reason": str}
        """
        import json

        from utils import prompts
        from utils.config.llm import create_llm

        preview = text[:500]
        prompt = prompts.DOCUMENT_RELEVANCE_CHECK_PROMPT.format(
            document_text_preview=preview,
        )
        try:
            llm = create_llm(label="문서적합성검증", temperature=0.0, max_tokens=256)
            response = llm.invoke([{"role": "user", "content": prompt}])
            raw = response.content.strip()
            # JSON 블록 추출
            if "```" in raw:
                raw = raw.split("```")[1]
                if raw.startswith("json"):
                    raw = raw[4:]
                raw = raw.strip()
            return json.loads(raw)
        except Exception:
            logger.warning("문서 적합성 검증 파싱 실패, 기본 허용 처리")
            return {"is_relevant": True, "category": "unknown", "reason": "검증 실패 — 기본 허용"}

    def _build_file_name(self, type_key: str, ext: str, company_name: str | None = None) -> str:
        """사용자 친화적 파일명을 생성합니다.

        형식: {회사이름}_{유형}_{날짜}_{순번:03d}.{ext}
        회사이름이 없는 경우: {유형}_{날짜}_{순번:03d}.{ext}
        """
        date_str = datetime.now().strftime("%Y%m%d")
        type_label = type_key.upper()
        existing = list(self.output_dir.glob(f"*_{type_label}_{date_str}_*.{ext}"))
        seq = len(existing) + 1
        # 기본값 sentinel을 None으로 정규화
        if company_name == "회사명 미기재":
            company_name = None
        if company_name:
            safe_name = company_name.replace("/", "_").strip()
            return f"{safe_name}_{type_label}_{date_str}_{seq:03d}.{ext}"
        return f"{type_label}_{date_str}_{seq:03d}.{ext}"

    def _build_docx_from_text(
        self,
        text: str,
        type_key: str,
        title: str,
        company_name: str | None = None,
    ) -> DocumentResponse:
        """LLM 생성 텍스트를 DOCX로 변환합니다.

        Args:
            text: LLM이 생성한 문서 텍스트
            type_key: 문서 유형 키
            title: 문서 제목

        Returns:
            문서 생성 응답
        """
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 제목
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 본문: 줄 단위로 파싱하여 제목/본문 구분
        for line in text.strip().split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()
                continue
            # "제N조" 패턴 → 소제목
            if stripped.startswith("제") and "조" in stripped[:6]:
                doc.add_heading(stripped, level=2)
            elif stripped.startswith("#"):
                # 마크다운 헤딩 제거
                clean = stripped.lstrip("#").strip()
                doc.add_heading(clean, level=2)
            else:
                doc.add_paragraph(stripped)

        # 파일 저장
        file_name = self._build_file_name(type_key, "docx", company_name)
        file_path = self.output_dir / file_name
        doc.save(file_path)

        with open(file_path, "rb") as f:
            file_content = base64.b64encode(f.read()).decode("utf-8")

        return DocumentResponse(
            success=True,
            doc_type_id=type_key,
            file_path=str(file_path),
            file_name=file_name,
            file_content=file_content,
            message=f"{title}이(가) 생성되었습니다.",
        )

    def _build_pdf_from_text(
        self,
        text: str,
        type_key: str,
        title: str,
        company_name: str | None = None,
    ) -> DocumentResponse:
        """LLM 생성 텍스트를 PDF로 변환합니다.

        Args:
            text: LLM이 생성한 문서 텍스트
            type_key: 문서 유형 키
            title: 문서 제목

        Returns:
            문서 생성 응답
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        _register_korean_font()

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "DocTitle",
            parent=styles["Title"],
            fontName="KoreanFont-Bold",
            fontSize=18,
            spaceAfter=30,
            alignment=1,
        )
        heading_style = ParagraphStyle(
            "DocHeading",
            parent=styles["Heading2"],
            fontName="KoreanFont-Bold",
            fontSize=13,
            spaceBefore=12,
            spaceAfter=6,
        )
        normal_style = ParagraphStyle(
            "DocNormal",
            parent=styles["Normal"],
            fontName="KoreanFont",
            fontSize=11,
            leading=16,
        )

        elements: list[Any] = []
        elements.append(Paragraph(xml_escape(title), title_style))
        elements.append(Spacer(1, 20))

        for line in text.strip().split("\n"):
            stripped = line.strip()
            if not stripped:
                elements.append(Spacer(1, 8))
                continue
            if stripped.startswith("제") and "조" in stripped[:6]:
                elements.append(Paragraph(f"<b>{xml_escape(stripped)}</b>", heading_style))
            elif stripped.startswith("#"):
                clean = stripped.lstrip("#").strip()
                elements.append(Paragraph(f"<b>{xml_escape(clean)}</b>", heading_style))
            else:
                elements.append(Paragraph(xml_escape(stripped), normal_style))

        doc.build(elements)

        file_name = self._build_file_name(type_key, "pdf", company_name)
        file_path = self.output_dir / file_name

        buffer.seek(0)
        with open(file_path, "wb") as f:
            f.write(buffer.read())

        buffer.seek(0)
        file_content = base64.b64encode(buffer.read()).decode("utf-8")

        return DocumentResponse(
            success=True,
            doc_type_id=type_key,
            file_path=str(file_path),
            file_name=file_name,
            file_content=file_content,
            message=f"{title}이(가) 생성되었습니다.",
        )

    def execute_action(
        self,
        action_type: str,
        params: dict[str, Any],
    ) -> dict[str, Any]:
        """액션을 실행합니다.

        Args:
            action_type: 액션 유형
            params: 액션 파라미터

        Returns:
            실행 결과
        """
        if action_type == "document_generation":
            doc_type = params.get("doc_type_id")
            if doc_type == "labor_contract":
                request = ContractRequest(**params)
                result = self.generate_labor_contract(request)
                return result.model_dump()

        return {"success": False, "message": f"지원하지 않는 액션: {action_type}"}
